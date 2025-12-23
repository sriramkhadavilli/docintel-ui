import os
from io import BytesIO

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from docx import Document

app = FastAPI()

# ✅ CORS (lets browser call this API)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # later restrict to your static web app domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DOCINTEL_ENDPOINT = os.environ.get("DOCINTEL_ENDPOINT")
DOCINTEL_KEY = os.environ.get("DOCINTEL_KEY")

client = None
if DOCINTEL_ENDPOINT and DOCINTEL_KEY:
    client = DocumentAnalysisClient(
        endpoint=DOCINTEL_ENDPOINT,
        credential=AzureKeyCredential(DOCINTEL_KEY)
    )

@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/convert")
async def convert(file: UploadFile = File(...)):
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="Only PDF is supported")

    if client is None:
        return JSONResponse(
            status_code=500,
            content={"error": "DOCINTEL_ENDPOINT / DOCINTEL_KEY not set in Container App env vars"}
        )

    pdf_bytes = await file.read()

    poller = client.begin_analyze_document("prebuilt-read", document=pdf_bytes)
    result = poller.result()

    doc = Document()
    doc.add_heading("Extracted Text", level=1)

    for page in result.pages:
        doc.add_heading(f"Page {page.page_number}", level=2)
        for line in page.lines:
            doc.add_paragraph(line.content)

    out = BytesIO()
    doc.save(out)
    out.seek(0)

    filename = (file.filename or "output.pdf").replace(".pdf", "") + ".docx"
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

